import tempfile

from docx import Document

from core.ats_check import check_ats_issues


def test_docx_with_table_flagged():
    doc = Document()
    doc.add_paragraph("Some resume text.")
    doc.add_table(rows=2, cols=2)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name

    issues = check_ats_issues(path, "Some resume text.")
    assert any("table" in i.lower() for i in issues)


def test_docx_without_table_is_clean():
    doc = Document()
    doc.add_paragraph("Some resume text with no tables at all.")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name

    issues = check_ats_issues(path, "Some resume text with no tables at all.")
    assert issues == []


def test_unsupported_extension_returns_no_issues():
    assert check_ats_issues("resume.txt", "plain text resume") == []
