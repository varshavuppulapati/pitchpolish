"""Builds a downloadable DOCX of the tailored resume bullets."""
import io


def build_docx(rewritten_bullets, job_title=None):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    heading = doc.add_heading("Tailored Resume Bullets", level=1)
    if job_title:
        sub = doc.add_paragraph(f"Tailored for: {job_title}")
        sub.runs[0].italic = True

    for item in rewritten_bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item["rewritten"])
        run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
