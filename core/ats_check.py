"""Lightweight, honest heuristics for common ATS (applicant tracking system) parsing traps.

These are not a full layout analysis - just the handful of signals that are
cheap to check and genuinely correlate with ATS parsing failures.
"""
import os


def check_ats_issues(file_path, resume_text):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _check_pdf(file_path, resume_text)
    if ext == ".docx":
        return _check_docx(file_path)
    return []


def _check_pdf(file_path, resume_text):
    from pypdf import PdfReader

    issues = []
    reader = PdfReader(file_path)
    page_count = len(reader.pages) or 1

    image_count = sum(len(page.images) for page in reader.pages)
    if image_count:
        issues.append("Contains embedded images — some ATS systems can't read text inside images.")

    chars_per_page = len(resume_text.strip()) / page_count
    if chars_per_page < 200:
        issues.append("Very little machine-readable text detected — this may be a scanned or image-based PDF that ATS can't parse at all.")

    lines = [ln.strip() for ln in resume_text.splitlines() if ln.strip()]
    if lines:
        short_lines = [ln for ln in lines if len(ln) < 25]
        if len(short_lines) / len(lines) > 0.45:
            issues.append("A lot of short, fragmented lines — often a sign of a multi-column layout, which ATS parsers tend to scramble.")

    return issues


def _check_docx(file_path):
    from docx import Document

    issues = []
    doc = Document(file_path)
    if doc.tables:
        issues.append(f"Contains {len(doc.tables)} table(s) — ATS systems often misread tabular layouts as jumbled text.")
    return issues
